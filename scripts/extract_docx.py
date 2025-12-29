from docx import Document
import sys
from pathlib import Path

if len(sys.argv) < 2:
    print("Usage: extract_docx.py <input.docx> [output.txt]")
    sys.exit(1)

input_path = Path(sys.argv[1])
output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path.with_suffix('.txt')

if not input_path.exists():
    print(f"File not found: {input_path}")
    sys.exit(2)

doc = Document(input_path)
lines = []

# Extract paragraphs
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        lines.append(text)

# Extract tables
for table in doc.tables:
    lines.append('\n[Table]')
    for row in table.rows:
        row_text = ' | '.join(cell.text.strip().replace('\n',' ') for cell in row.cells)
        lines.append(row_text)

output_path.write_text('\n'.join(lines), encoding='utf-8')
print(f"Extracted text written to: {output_path}")